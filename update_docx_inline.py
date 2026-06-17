# -*- coding: utf-8 -*-
import docx

doc = docx.Document('Frontend_System_Design_and_Implementation_Final_Normal.docx')

# First, clean up previous insertions
paragraphs = doc.paragraphs
paras_to_remove = []

for i, p in enumerate(paragraphs):
    if "Code Implementation" in p.text or "Brief Explanation" in p.text:
        paras_to_remove.append(p)
        # Also remove the next paragraph which contains the actual code/desc
        if i + 1 < len(paragraphs):
            paras_to_remove.append(paragraphs[i+1])

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        p._p = p._element = None

for p in paras_to_remove:
    try:
        delete_paragraph(p)
    except:
        pass

# Now define all snippets to insert
code_snippets = {
    '4.1 Notification Management System': {
        'code': '''Future<void> fetchNotifications() async {
  try {
    emit(NotificationLoading());
    final response = await dio.get(EndPoints.getNotifications);
    final notifications = (response.data as List).map((n) => NotificationModel.fromJson(n)).toList();
    emit(NotificationSuccess(notifications));
  } catch (e) {
    emit(NotificationError(e.toString()));
  }
}''',
        'desc': 'This code fetches real-time alerts from the backend. The Cubit emits a Loading state to show a spinner, then makes an authenticated GET request. If successful, it parses the JSON into Dart models and updates the UI.'
    },
    '4.21 File Upload and Streaming': {
        'code': '''static Future<String?> viewFileDirectly({
  required String filePath, required String fileName
}) async {
  try {
    final tempDir = await getTemporaryDirectory();
    final tempPath = '${tempDir.path}/$fileName';
    await _dio.download(url, tempPath);
    if (Platform.isWindows) { await Process.run('explorer', [tempPath]); }
    else { await OpenFilex.open(tempPath); }
    return null;
  } catch (e) {
    return 'فشل فتح الملف: ${e.message}';
  }
}''',
        'desc': 'This handles secure file streaming. It downloads the evidence file using Dio (which includes the JWT token) into a temporary directory, then seamlessly opens it in the native OS viewer without requiring manual saving.'
    },
    '4.22 Authentication and Interceptors': {
        'code': '''class LoginInterceptor extends Interceptor {
  @override
  void onRequest(RequestOptions options, RequestInterceptorHandler handler) async {
    final token = await LoginStorage.getToken();
    if (token != null) {
      options.headers['Authorization'] = 'Bearer $token';
    }
    super.onRequest(options, handler);
  }
}''',
        'desc': 'This Dio Interceptor is the core of our authentication. It automatically intercepts every outgoing HTTP request across the entire application and safely injects the JWT Bearer token into the headers, ensuring secure access.'
    },
    '5. Spotlight Feature: AI Course Description Generation Engine': {
        'code': '''Future<void> _pollGenerationStatus() async {
  bool isFinished = false;
  while (!isFinished) {
    final statusResult = await AiDescriptionService.checkGenerationStatus(generationId: id);
    if (statusResult.data?.status == 'Generating') {
      await Future.delayed(const Duration(seconds: 10)); // Yield thread
    } else {
      isFinished = true;
    }
  }
  isGenerationCompleted = true;
}''',
        'desc': 'This is the long-polling mechanism. It queries the AI microservice every 10 seconds. The `await Future.delayed` is crucial because it yields the main thread, keeping the UI fully responsive while waiting for the heavy AI processing.'
    },
    '4.2 Administrative Dashboard': {
        'code': '''Widget buildIndicatorsChart(List<IndicatorData> data) {
  return SfCartesianChart(
    primaryXAxis: CategoryAxis(),
    series: <ChartSeries>[
      ColumnSeries<IndicatorData, String>(
        dataSource: data,
        xValueMapper: (IndicatorData ind, _) => ind.name,
        yValueMapper: (IndicatorData ind, _) => ind.progress,
        color: AppColors.primary,
      )
    ]
  );
}''',
        'desc': 'This builds the interactive visual charts on the Admin Dashboard using the Syncfusion library. It maps the dynamic indicator data fetched from the Cubit directly into a responsive, color-coded Column Series.'
    },
    '4.14 Task Assignment Engine': {
        'code': '''Future<void> assignTask(int doctorId, int indicatorId) async {
  try {
    emit(TaskAssignLoading());
    await dio.post(EndPoints.assignIndicator, data: {
      'doctorId': doctorId,
      'indicatorId': indicatorId,
      'deadline': deadlineController.text,
    });
    emit(TaskAssignSuccess());
  } catch (e) {
    emit(TaskAssignError('Failed to assign task'));
  }
}''',
        'desc': 'This core workflow engine code assigns an indicator to a specific doctor with a hard deadline. It updates the database via POST request, triggering a state change that reflects immediately on the UI.'
    },
    '4.4 Settings and Profile Management': {
        'code': '''void toggleLanguage() {
  if (context.locale.languageCode == 'ar') {
    context.setLocale(const Locale('en'));
  } else {
    context.setLocale(const Locale('ar'));
  }
  emit(LanguageChangedState());
}''',
        'desc': 'This code allows users to toggle the entire application layout natively. By changing the locale, the Flutter framework automatically mirrors the structure (RTL for Arabic, LTR for English) across all screens instantly.'
    }
}

# Iterate through document to find insertion points
# We want to insert AFTER the "Workflow & State Management:" explanation paragraph.
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    
    # Check section headers and insert code in the NEXT paragraph (which is usually the workflow explanation)
    for section_title, snippet in code_snippets.items():
        if section_title in text:
            # We found the section! Now, look ahead for the Workflow paragraph
            for j in range(i+1, min(i+10, len(doc.paragraphs))):
                if "Workflow & State Management" in doc.paragraphs[j].text:
                    # The paragraph after "Workflow & State Management:" is the actual explanation.
                    # We want to insert AFTER that explanation.
                    target_p = doc.paragraphs[j+1]
                    
                    # Insert Code
                    new_p = target_p.insert_paragraph_before("")
                    new_p.add_run("\n💻 كود برمجي (Code Implementation):\n").bold = True
                    run = new_p.add_run(snippet['code'])
                    run.font.name = 'Courier New'
                    
                    # Insert Explanation
                    new_p2 = target_p.insert_paragraph_before("")
                    new_p2.add_run("\n💡 شرح الكود باختصار:\n").bold = True
                    new_p2.add_run(snippet['desc'] + "\n")
                    break

# Spotlight feature has a different structure
for i, paragraph in enumerate(doc.paragraphs):
    if "Below is a critical excerpt from the system demonstrating how the Cubit autonomously polls" in paragraph.text:
        target_p = doc.paragraphs[i]
        new_p = target_p.insert_paragraph_before("")
        new_p.add_run("\n💻 كود برمجي (Code Implementation):\n").bold = True
        run = new_p.add_run(code_snippets['5. Spotlight Feature: AI Course Description Generation Engine']['code'])
        run.font.name = 'Courier New'
        
        new_p2 = target_p.insert_paragraph_before("")
        new_p2.add_run("\n💡 شرح الكود باختصار:\n").bold = True
        new_p2.add_run(code_snippets['5. Spotlight Feature: AI Course Description Generation Engine']['desc'] + "\n")


doc.save('Frontend_System_Design_and_Implementation_Final_Normal.docx')
print('Successfully cleaned and added inline code snippets!')
