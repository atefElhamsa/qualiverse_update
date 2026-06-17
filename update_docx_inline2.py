# -*- coding: utf-8 -*-
import docx

doc = docx.Document('Frontend_System_Design_and_Implementation_Final_Normal.docx')

code_snippets = {
    '4.3 Faculty Dashboard Overview': {
        'code': '''Future<void> fetchDoctorTasks() async {
  try {
    emit(DoctorTasksLoading());
    final response = await dio.get(EndPoints.getDoctorTasks);
    final tasks = (response.data as List).map((t) => TaskModel.fromJson(t)).toList();
    emit(DoctorTasksSuccess(tasks));
  } catch (e) {
    emit(DoctorTasksError(e.toString()));
  }
}''',
        'desc': 'This fetches the personalized dashboard data for a faculty member. It retrieves only the indicators assigned specifically to their ID and their respective deadlines, ensuring data privacy and a focused user experience.'
    },
    '4.5 Course Directory and Management': {
        'code': '''Future<void> createCourse(CourseModel course) async {
  try {
    emit(CreateCourseLoading());
    await dio.post(EndPoints.createCourse, data: course.toJson());
    emit(CreateCourseSuccess());
  } catch (e) {
    emit(CreateCourseError(e.toString()));
  }
}''',
        'desc': 'This function sends a POST request with the new course payload (including credit hours and prerequisites) to the backend. It uses strongly-typed Dart objects and serializes them to JSON seamlessly.'
    },
    '4.8 AI Course Description Wizard': {
        'code': '''Future<void> startAiGeneration({required int courseId}) async {
  emit(AiDescriptionStartLoading());
  try {
    final result = await AiDescriptionService.startGeneration(courseId: courseId);
    generationId = result.data?.id;
    emit(AiDescriptionStartSuccess(generationId!));
  } catch (e) {
    emit(AiDescriptionStartError(e.toString()));
  }
}''',
        'desc': 'This triggers the initial AI document generation workflow on the backend microservice. It securely receives a unique `generationId` which is then used by the polling engine to track the progress of the PDF extraction.'
    },
    '4.11 Cycle Tabs and Navigation': {
        'code': '''class AppRouter {
  static Route? generateRoute(RouteSettings settings) {
    switch (settings.name) {
      case Routes.cycleTabsRoute:
        return MaterialPageRoute(builder: (_) => const CycleTabsScreen());
      case Routes.homeRoute:
        return MaterialPageRoute(builder: (_) => const HomeScreen());
      default:
        return MaterialPageRoute(builder: (_) => const NotFoundScreen());
    }
  }
}''',
        'desc': 'This code demonstrates the central routing architecture. It uses Flutter’s native Navigator 2.0 concept to manage complex nested views, allowing users to switch between cycle tabs without losing their navigation history.'
    },
    '4.16 Approval and Rejection Workflows': {
        'code': '''Future<void> rejectEvidence(int evidenceId, String reason) async {
  try {
    emit(ApprovalLoading());
    await dio.post(EndPoints.rejectEvidence(evidenceId), data: {'reason': reason});
    emit(ApprovalRejectedSuccess());
  } catch (e) {
    emit(ApprovalError('Rejection failed'));
  }
}''',
        'desc': 'This allows the Quality Assurance admin to reject a submitted file. It requires an explicit reasoning string to be passed to the backend so the doctor knows exactly why their evidence was not accepted.'
    },
    '4.24 Offline Connectivity Fallback': {
        'code': '''Future<bool> checkInternet() async {
  final connectivityResult = await Connectivity().checkConnectivity();
  if (connectivityResult == ConnectivityResult.none) {
    emit(NoInternetState());
    return false;
  }
  return true;
}''',
        'desc': 'This is our defensive programming in action. Before ANY critical API request is fired, this method checks the physical hardware connection. If the user is offline, it blocks the request and shows a graceful "No Internet" screen instead of throwing unhandled exceptions.'
    },
    '4.25 Export and Reporting Engine': {
        'code': '''Future<void> exportReport(int programId) async {
  try {
    emit(ExportLoading());
    final response = await dio.get(
      EndPoints.exportProgramReport(programId),
      options: Options(responseType: ResponseType.bytes),
    );
    await FileHelper.saveFile(response.data, 'Report_$programId.pdf');
    emit(ExportSuccess());
  } catch (e) {
    emit(ExportError(e.toString()));
  }
}''',
        'desc': 'This handles massive binary data fetching. By setting `ResponseType.bytes`, Dio knows not to parse the incoming PDF as a JSON string, preventing memory corruption and ensuring the file is saved cleanly to the disk.'
    },
    '4.9 User and Role Management': {
        'code': '''Future<void> assignRole(String userId, String roleName) async {
  try {
    emit(RoleAssignLoading());
    await dio.post(EndPoints.assignRole, data: {
      'userId': userId,
      'role': roleName,
    });
    emit(RoleAssignSuccess());
  } catch (e) {
    emit(RoleAssignError(e.toString()));
  }
}''',
        'desc': 'This administrative endpoint modifies the JWT claims for a specific user dynamically. Once the role is updated, the target user will receive new permissions upon their next token refresh cycle.'
    },
    '4.6 Digital Course Folders': {
        'code': '''Future<void> getFolderFiles(int folderId) async {
  try {
    emit(FolderFilesLoading());
    final response = await dio.get(EndPoints.getFolderFiles(folderId));
    final files = (response.data as List).map((f) => FileModel.fromJson(f)).toList();
    emit(FolderFilesSuccess(files));
  } catch (e) {
    emit(FolderFilesError(e.toString()));
  }
}''',
        'desc': 'This retrieves the physical layout of a Quality Assurance folder. It maps the incoming list of files and subdirectories into a tree-like Dart structure `FileModel` which the UI recursively renders.'
    },
    '4.10 Academic Year Configuration': {
        'code': '''Future<void> archiveAcademicYear(int yearId) async {
  try {
    emit(ArchiveYearLoading());
    await dio.put(EndPoints.archiveYear(yearId));
    emit(ArchiveYearSuccess());
  } catch (e) {
    emit(ArchiveYearError(e.toString()));
  }
}''',
        'desc': 'This triggers a soft-delete mechanism on the backend. Archiving a year prevents it from appearing in active dropdowns while preserving all historical indicator data for accreditation history and long-term charts.'
    }
}

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    
    for section_title, snippet in code_snippets.items():
        if section_title in text:
            for j in range(i+1, min(i+10, len(doc.paragraphs))):
                if "Workflow & State Management" in doc.paragraphs[j].text:
                    # Check if code is already inserted here to avoid duplicates
                    if j+2 < len(doc.paragraphs) and "كود برمجي" in doc.paragraphs[j+2].text:
                        break
                        
                    target_p = doc.paragraphs[j+1]
                    
                    new_p = target_p.insert_paragraph_before("")
                    new_p.add_run("\n💻 كود برمجي (Code Implementation):\n").bold = True
                    run = new_p.add_run(snippet['code'])
                    run.font.name = 'Courier New'
                    
                    new_p2 = target_p.insert_paragraph_before("")
                    new_p2.add_run("\n💡 شرح الكود باختصار:\n").bold = True
                    new_p2.add_run(snippet['desc'] + "\n")
                    break

doc.save('Frontend_System_Design_and_Implementation_Final_Normal.docx')
print('Successfully added MORE inline code snippets!')
