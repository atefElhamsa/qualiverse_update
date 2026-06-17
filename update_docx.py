import docx

doc_path = 'Frontend_System_Design_and_Implementation_Final_Normal.docx'
doc = docx.Document(doc_path)

replacements = {
    "Future<void> fetchNotifications() async {": """  Future<void> getAllNotifications({
    bool loadMore = false,
    bool? isRead,
  }) async {
    if (loadMore) {
      if (!hasMore) return;
      currentPage++;
      emit(NotificationsPaginationLoading());
    } else {
      currentPage = 1;
      allNotifications.clear();
      hasMore = true;
      currentIsReadFilter = isRead;
      emit(NotificationsLoading());
    }

    final result = await NotificationsService.getNotifications(
      pageIndex: currentPage,
      pageSize: pageSize,
      isRead: currentIsReadFilter,
    );

    if (result.isSuccess) {
      final List<NotificationModel> newData = result.data ?? [];
      if (newData.length < pageSize) {
        hasMore = false;
      }
      allNotifications.addAll(newData);
      emit(NotificationsSuccess(notifications: List.from(allNotifications), hasMore: hasMore));
    } else {
      emit(NotificationsError(result.error?.description ?? "Error"));
    }
  }""",
    
    "Future<void> fetchDoctorTasks() async {": """  Future<void> fetchAssignments({
    required int academicYearId,
    String? doctorId,
    String? status,
  }) async {
    emit(AssignmentsLoading());
    try {
      final response = await AssignmentIndicatorAdminService.getAssignmentIndicatorsAdmin(
            academicYearId: academicYearId,
            doctorId: doctorId,
            status: status,
          );
      assignments = response.data!;
      emit(AssignmentsLoaded(assignments: assignments));
    } catch (e) {
      emit(AssignmentsError(error: e.toString()));
    }
  }""",
    
    "Future<void> createCourse(CourseModel course) async {": """  Future<void> fetchCourses({
    required int yearId,
    required int levelId,
    required int semesterId,
    int? departmentId,
    String? lang,
  }) async {
    emit(CourseLoading());
    try {
      final data = await CourseService.getCourses(
        yearId: yearId,
        levelId: levelId,
        termId: semesterId,
        departmentId: departmentId,
        lang: lang,
      );
      if (isClosed) return;
      courses = data.courses!;
      emit(CourseSuccess(courses: courses, selectedCourse: selectedCourse));
    } catch (e) {
      if (isClosed) return;
      final msg = e.toString().replaceFirst('Exception: ', '').trim();
      if (msg.contains('No Internet')) {
        emit(CourseError(message: 'Check your internet connection'));
      } else {
        emit(CourseError(message: msg));
      }
    }
  }""",
    
    "Future<void> getFolderFiles(int folderId) async {": """  Future<void> fetchCourseFolders({required int courseId}) async {
    currentCourseId = courseId;
    emit(CourseFolderLoading());
    try {
      final data = await CourseFolderService.getCourseFolders(courseId: courseId);
      courseFolders = data.courseFolders!;
      emit(CourseFolderSuccess(courseFolders: courseFolders, selectedCourseFolder: selectedCourseFolder));
    } catch (e) {
      emit(CourseFolderError(message: 'Something went wrong'));
    }
  }""",
    
    "Future<void> startAiGeneration({required int courseId}) async {": """  Future<void> startAiGeneration({required int courseId}) async {
    emit(AiDescriptionStartLoading());
    try {
      final result = await AiDescriptionService.startGeneration(courseId: courseId);
      generationId = result.data?.id;
      emit(AiDescriptionStartSuccess(generationId!));
    } catch (e) {
      emit(AiDescriptionStartError(e.toString().replaceFirst('Exception: ', '').trim()));
    }
  }""",
    
    "static Future<String?> viewFileDirectly({": """  static Future<String?> viewFileDirectly({
    required String filePath,
    required String fileName,
    void Function(int received, int total)? onProgress,
  }) async {
    try {
      final List<ConnectivityResult> connectivityResult = await (Connectivity().checkConnectivity());
      if (connectivityResult.contains(ConnectivityResult.none)) {
        return 'النت قطع، يرجى التحقق من اتصالك بالإنترنت';
      }

      final url = filePath.startsWith('http')
          ? filePath
          : '${EndPoints.baseUrlToOpenFile}/${filePath.startsWith('/') ? filePath.substring(1) : filePath}';

      final tempDir = await getTemporaryDirectory();
      final tempPath = '${tempDir.path}/$fileName';

      await _dio.download(
        url,
        tempPath,
        onReceiveProgress: onProgress,
        deleteOnError: true,
        options: Options(
          responseType: ResponseType.bytes,
          receiveTimeout: const Duration(seconds: 30),
          headers: {'Accept': '*/*'},
        ),
      );

      if (Platform.isWindows) {
        await Process.run('explorer', [tempPath]);
      } else {
        await OpenFilex.open(tempPath);
      }
      return null;""",
      
      "class LoginInterceptor extends Interceptor {": """  @override
  Future<void> onError(DioException err, ErrorInterceptorHandler handler) async {
    if (loggedOut || isAuthPath(err.requestOptions.path)) {
      handler.next(err);
      return;
    }

    if (err.response?.statusCode == 401) {
      if (isRefreshing) {
        final completer = Completer<Response>();
        queue.add(QueuedRequest(err.requestOptions, completer));
        try {
          final response = await completer.future;
          handler.resolve(response);
        } catch (e) {
          handler.next(e is DioException ? e : err);
        }
        return;
      }

      isRefreshing = true;
      final refreshed = await refreshToken();
      isRefreshing = false;

      if (!refreshed) {
        failAllQueued(err);
        logout();
        handler.next(err);
        return;
      }
      // Retry logic omitted for brevity...
  }""",
  
      "Future<void> _pollGenerationStatus() async {": """  Future<void> _pollGenerationStatus() async {
    try {
      bool isFinished = false;
      while (!isFinished) {
        final statusResult = await AiDescriptionService.checkGenerationStatus(
          generationId: generationId!,
        );
        final status = statusResult.data?.status;

        if (status == 'Generating') {
          await Future.delayed(const Duration(seconds: 10)); // Yield thread
        } else if (status == 'Failed' || status == 'Error') {
          throw Exception(statusResult.data?.error ?? 'Generation failed');
        } else if (status == 'Complete') {
          isFinished = true;
        } else {
          isFinished = true;
        }
      }
      isGenerationCompleted = true;
    } catch (e) {
      isGenerationCompleted = false;
    }
  }"""
}

# Iterate and replace
for i, p in enumerate(doc.paragraphs):
    for key, new_code in replacements.items():
        if key in p.text:
            # We found the start of a fake code block. Replace it.
            p.text = new_code
            
            # Now we need to clear out the rest of the fake code block in subsequent paragraphs
            # until we hit an empty line or normal text
            j = i + 1
            while j < len(doc.paragraphs):
                next_p = doc.paragraphs[j].text.strip()
                if next_p == "" or next_p.startswith("This handles") or next_p.startswith("When a user") or next_p.startswith("This is the long-polling") or next_p.startswith("Architectural Significance"):
                    break
                doc.paragraphs[j].text = "" # Clear the fake code lines
                j += 1

doc.save('Frontend_System_Design_and_Implementation_Final_Normal.docx')
print('Successfully updated docx')
